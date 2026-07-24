#!/usr/bin/env python3
"""生成《钉钉机器人功能介绍》Word 文档。"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "钉钉机器人功能介绍.docx"

# 钉钉品牌蓝 + 辅助色
COLOR_PRIMARY = RGBColor(0x00, 0x89, 0xFF)
COLOR_ACCENT = RGBColor(0x1A, 0x73, 0xE8)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_HEADER_BG = "E8F4FD"
COLOR_ALT_ROW = "F7FBFF"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def style_heading(paragraph, level: int = 1) -> None:
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    if not paragraph.runs:
        pass
    for r in paragraph.runs:
        r.font.color.rgb = COLOR_PRIMARY if level == 1 else COLOR_ACCENT
        r.font.bold = True
        sizes = {1: 18, 2: 14, 3: 12}
        r.font.size = Pt(sizes.get(level, 12))
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_body(doc: Document, text: str, *, bold: bool = False, color=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.color.rgb = color or COLOR_TEXT
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], COLOR_HEADER_BG)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_ACCENT
                r.font.name = "PingFang SC"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = COLOR_ALT_ROW if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = val
            set_cell_shading(cells[ci], bg)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "PingFang SC"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
                    r.font.color.rgb = COLOR_TEXT
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # —— 封面 ——
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Yaahlan 智能工具\n钉钉机器人")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = COLOR_PRIMARY
    tr.font.name = "PingFang SC"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("功能介绍与使用指南")
    sr.font.size = Pt(16)
    sr.font.color.rgb = COLOR_MUTED
    sr.font.name = "PingFang SC"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"更新日期：{date.today().isoformat()}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # —— 1 简介 ——
    h = doc.add_heading("一、产品简介", level=1)
    style_heading(h, 1)
    add_body(
        doc,
        "Yaahlan 智能工具钉钉机器人，是连接钉钉群与 Cursor AI 智能体的「无人值守网关」。"
        "团队成员在群里 @机器人 即可查询测试数据、生成测试用例、导出钉钉文档、"
        "执行 MOA/Admin 常用操作——无需打开 IDE，也无需手动点 Run。",
    )
    add_body(
        doc,
        "机器人背后是一台常驻运行的执行机（Mac），通过 Cursor Agent 自动调用仓库内各模块脚本与钉钉 MCP，"
        "把结果以自然语言、表格或在线文档链接的形式回传到群里。",
    )

    # —— 2 快速上手 ——
    h = doc.add_heading("二、30 秒快速上手", level=1)
    style_heading(h, 1)
    add_table(
        doc,
        ["步骤", "操作"],
        [
            ["1", "在已添加企业机器人的钉钉群里，输入 @机器人"],
            ["2", "附上你的问题或任务（见下文示例）"],
            ["3", "机器人先回复「已收到，执行中…」，完成后回传结果"],
        ],
    )
    add_body(doc, "最简示例：", bold=True)
    add_bullet(doc, "@机器人 查一下用户 100465989 的 VIP 等级")

    # —— 3 两种交互方式 ——
    h = doc.add_heading("三、两种交互方式", level=1)
    style_heading(h, 1)
    add_table(
        doc,
        ["方式", "特点", "适用场景"],
        [
            [
                "快捷指令",
                "不走 AI，秒级响应",
                "环境检查、MOA 探活、固定格式导出、VIP 升级、生成测试报告等",
            ],
            [
                "自然语言",
                "走 Cursor Agent，支持多轮上下文",
                "查数、抓包、读钉钉文档、生成用例、复杂组合任务",
            ],
        ],
    )
    add_body(
        doc,
        "提示：执行机本地可运行 python3 platform/open_catalog.py 打开工具能力目录。",
        color=COLOR_MUTED,
    )

    # —— 4 快捷指令 ——
    h = doc.add_heading("四、快捷指令一览", level=1)
    style_heading(h, 1)
    add_body(doc, "以下指令无需复杂描述，直接 @机器人 发送即可：")
    add_table(
        doc,
        ["指令", "说明"],
        [
            ["环境检查", "本机配置自检（Bridge、凭证、依赖等）"],
            ["MOA检查", "测试环境 MOA Cookie 是否有效"],
            ["导出 temporary_testcase/xxx.csv", "将本地用例文件导出为钉钉在线表格"],
            ["2.4.5版本生成测试报告", "生成指定版本内/外网 HTML 测试报告（zip 附件）"],
            ["100465989升级 VIP3", "MOA VIP 升级（用户 ID + 等级）"],
            ["中断操作", "打断本群当前正在执行的任务"],
            ["重新执行", "重跑本群上一条任务"],
        ],
    )

    # —— 5 智能能力 ——
    h = doc.add_heading("五、智能对话能做什么", level=1)
    style_heading(h, 1)
    add_body(doc, "通过自然语言，机器人可调用工具平台已登记的各模块能力，主要包括：")
    add_table(
        doc,
        ["模块", "典型能力"],
        [
            ["Admin 后台", "查用户详情、设备信息、在线状态等"],
            ["MOA", "VIP/贵族、房间、礼物背包、公会、家族、道具、风控等 Stage 接口"],
            ["Gift Stage 送礼", "HTTP /v2/gift/send 房间内/私聊/群组送礼"],
            ["Tunnel 抓包", "只读查询用户 HTTP 请求与响应"],
            ["Risk 风控", "解除设备/手机号风控等"],
            ["钉钉文档", "读取 alidocs 需求文档、同步用例到 Excel"],
            ["线上环境", "消息含「线上环境」时走 online/ 统一入口（Admin/MOA/Tunnel）"],
        ],
    )

    h2 = doc.add_heading("典型场景示例", level=2)
    style_heading(h2, 2)
    add_table(
        doc,
        ["你想做的事", "可以这样 @机器人"],
        [
            ["了解仓库工具有哪些", "介绍一下 platform 目录有哪些工具"],
            ["查用户信息", "查一下用户 100465989 的详情"],
            ["抓包验证接口", "查用户 100465989 最近 1 小时的送礼接口"],
            ["测试环境送礼", "给用户 100465989 在房间 xxx 送玫瑰礼物"],
            ["读需求文档", "读这个文档 https://alidocs.dingtalk.com/..."],
            ["生成测试用例", "根据这份 PRD 生成测试用例"],
            ["附图分析", "（发截图）看这个页面有什么问题"],
        ],
    )

    # —— 6 结果展示 ——
    h = doc.add_heading("六、结果怎么展示", level=1)
    style_heading(h, 1)
    add_bullet(doc, "查询类结果默认直接在群里展示，优先用 Markdown 表格，语言自然易懂")
    add_bullet(doc, "用户列表默认只展示前 10 条；需要全量时说「查看全部数据」")
    add_bullet(doc, "需要写入钉钉文档时说「导出到钉钉文档」——成功时群里只回在线表格链接")
    add_bullet(doc, "生成测试用例会自动写入 temporary_testcase/ 并同步到钉钉文档")
    add_bullet(doc, "大段表格/JSON 也会自动导出到「智能工具导出」目录并回链接")

    # —— 7 任务控制 ——
    h = doc.add_heading("七、任务控制与排队", level=1)
    style_heading(h, 1)
    add_bullet(doc, "收到消息后会先回复「已收到，执行中…」；长任务会定期发送心跳（已执行时长）")
    add_bullet(doc, "本群同时只跑一个任务；若排队会提示前面约几个任务、预计等待时间")
    add_bullet(doc, "执行中可随时发「中断操作」打断当前任务")
    add_bullet(doc, "发「重新执行」可重跑上一条任务（无需重新输入）")
    add_bullet(doc, "每人有独立 Cursor 会话窗口，多轮对话自动带上下文")

    # —— 8 限制 ——
    h = doc.add_heading("八、重要限制（请务必了解）", level=1)
    style_heading(h, 1)
    add_table(
        doc,
        ["限制项", "说明"],
        [
            ["默认测试环境", "未写「线上环境」时，只用 Admin/MOA/Tunnel 测试环境，不会碰线上"],
            ["不支持真机 ADB", "macro、observe、capture、礼物面板 UI 点按等须在 Cursor 本机执行"],
            ["送礼默认 HTTP", "说「送礼」默认走 Gift HTTP；仅明确「背包送礼」时才走 MOA 背包"],
            ["抓包只读", "Tunnel 仅查询，不会在钉钉侧发起真机操作"],
            ["代码修改权限", "仅白名单账号可通过机器人改网关/Agent 代码；其他人只读模式"],
            ["执行机需在线", "执行机休眠或网关未运行时，@ 机器人会无响应"],
        ],
    )

    # —— 9 FAQ ——
    h = doc.add_heading("九、常见问题", level=1)
    style_heading(h, 1)
    faqs = [
        (
            "Q：@ 了机器人没反应？",
            "A：检查执行机是否唤醒、网关是否在运行。管理员可执行 ./gateway_ctl.sh status 查看状态。",
        ),
        (
            "Q：怎么查完整用户列表？",
            "A：在首次只展示 10 条后，回复「查看全部数据」即可。",
        ),
        (
            "Q：测试用例生成后在哪看？",
            "A：机器人会自动同步到钉钉在线表格并在群里回链接；本地文件在 temporary_testcase/。",
        ),
        (
            "Q：能做真机自动化吗？",
            "A：不能。请在 Cursor IDE 本机使用 ADB 相关能力。",
        ),
        (
            "Q：如何获取全部工具能力清单？",
            "A：执行机本地可运行 python3 platform/open_catalog.py。",
        ),
    ]
    for q, a in faqs:
        add_body(doc, q, bold=True)
        add_body(doc, a)

    # —— 附录 ——
    h = doc.add_heading("附录：群内一句话备忘", level=1)
    style_heading(h, 1)
    add_body(
        doc,
        "@机器人 + 你的需求 ｜ MOA检查 ｜ 中断操作 ｜ 重新执行 ｜ "
        "查看全部数据 ｜ 导出到钉钉文档",
        color=COLOR_ACCENT,
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    main()
